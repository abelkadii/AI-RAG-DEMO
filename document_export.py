"""Export helpers for generated document reports."""

from __future__ import annotations

from io import BytesIO
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from document_models import DocumentTrace
import pymupdf


def markdown_bytes(trace: DocumentTrace) -> bytes:
    return trace.final_markdown.encode("utf-8")


def docx_bytes(trace: DocumentTrace) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Title"].font.name = "Aptos Display"
    styles["Title"].font.size = Pt(24)

    title = document.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(trace.plan.title)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Knowledge Base: {trace.spec.knowledge_base}")

    document.add_paragraph(f"Audience: {trace.spec.audience}")
    document.add_paragraph(f"Client Brief: {trace.spec.client_brief}")

    for section_trace in trace.sections:
        if section_trace.section_id == "evidence":
            continue
        document.add_heading(section_trace.title, level=1)
        for block in section_trace.content_markdown.splitlines():
            text = block.strip()
            if not text:
                continue
            if text.startswith("- "):
                document.add_paragraph(text[2:], style="List Bullet")
            elif re_numbered(text):
                document.add_paragraph(text[text.find(" ") + 1 :], style="List Number")
            else:
                document.add_paragraph(text)

    references = sorted(set(re.findall(r"\[([^\[\]\n]+? p\.\d+)\]", trace.final_markdown)))
    if references:
        document.add_heading("Evidence / References", level=1)
        for reference in references:
            document.add_paragraph(f"[{reference}]", style="List Bullet")

    document.add_heading("Quality Review", level=1)
    document.add_paragraph(trace.final_qc.summary)
    document.add_paragraph(f"Citation validation: {'passed' if trace.citation_validation.valid else 'needs review'}")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Evidence-grounded AI document production demo")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def pdf_bytes(trace: DocumentTrace) -> bytes:
    """Render a lightweight, dependency-free PDF from the final Markdown."""
    pdf = pymupdf.open()
    page = pdf.new_page(width=595, height=842)
    margin = 48
    y = margin
    for raw in trace.final_markdown.splitlines():
        text = raw.strip()
        if not text:
            y += 8
            continue
        if text.startswith("# "):
            size = 20
            text = text[2:]
        elif text.startswith("## "):
            size = 14
            text = text[3:]
        else:
            size = 9.5
            if text.startswith("- "):
                text = "• " + text[2:]
        height = max(18, size * 1.55)
        rect = pymupdf.Rect(margin, y, 595 - margin, y + height * 3)
        used = page.insert_textbox(rect, text, fontsize=size, fontname="helv", color=(0.08, 0.1, 0.14))
        y += max(height, abs(used) if used < 0 else height)
        if y > 790:
            page = pdf.new_page(width=595, height=842)
            y = margin
    return pdf.tobytes()


def re_numbered(text: str) -> bool:
    prefix = text.split(" ", 1)[0]
    return prefix.endswith(".") and prefix[:-1].isdigit()
